"""Genera los documentos Word del proyecto: general-técnico y publicitario.

Se escriben con python-docx directamente sobre el contenido real del repositorio
para que no queden desactualizados respecto del código.
"""

from __future__ import annotations

import datetime as dt
import pathlib

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

AZUL = RGBColor(0x04, 0x07, 0x64)
AZUL_SEC = RGBColor(0x1C, 0x73, 0xCB)
GRIS = RGBColor(0x54, 0x54, 0x54)

DOCS = pathlib.Path("docs")
CAPTURAS = DOCS / "capturas"
FECHA = dt.date(2026, 8, 18)
VERSION = "1.2.0"


def preparar(documento: Document) -> None:
    normal = documento.styles["Normal"]
    normal.font.name = "Roboto"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor(0x35, 0x35, 0x35)
    for nivel, tamano in ((1, 20), (2, 15), (3, 12)):
        estilo = documento.styles[f"Heading {nivel}"]
        estilo.font.name = "Roboto"
        estilo.font.size = Pt(tamano)
        estilo.font.color.rgb = AZUL
        estilo.font.bold = True


def portada(documento: Document, titulo: str, bajada: str) -> None:
    logo = DOCS / "marca" / "logo_flash_campaign.png"
    if logo.exists():
        parrafo = documento.add_paragraph()
        parrafo.add_run().add_picture(str(logo), width=Inches(2.6))

    encabezado = documento.add_paragraph()
    corrida = encabezado.add_run(titulo)
    corrida.font.size = Pt(24)
    corrida.font.bold = True
    corrida.font.color.rgb = AZUL

    sub = documento.add_paragraph()
    corrida = sub.add_run(bajada)
    corrida.font.size = Pt(12)
    corrida.font.color.rgb = AZUL_SEC

    ficha = documento.add_paragraph()
    corrida = ficha.add_run(
        f"Proyecto: Flash Campaign · Nombre normalizado: flash_campaign\n"
        f"Versión: {VERSION} · Estado: entregado para revisión\n"
        f"Fecha: {FECHA.strftime('%d-%m-%Y')} · Autor: Gamonal\n"
        f"Repositorio: https://github.com/andresgamonalm/flash-campaign\n"
        f"Subdominio previsto: https://flash-campaign.gamonal.app"
    )
    corrida.font.size = Pt(9.5)
    corrida.font.color.rgb = GRIS
    documento.add_page_break()


def vinetas(documento: Document, items: list[str]) -> None:
    for item in items:
        documento.add_paragraph(item, style="List Bullet")


def tabla(documento: Document, encabezados: list[str], filas: list[list[str]]) -> None:
    t = documento.add_table(rows=1, cols=len(encabezados))
    t.style = "Light Grid Accent 1"
    for celda, texto in zip(t.rows[0].cells, encabezados):
        celda.text = ""
        corrida = celda.paragraphs[0].add_run(texto)
        corrida.font.bold = True
        corrida.font.size = Pt(9.5)
    for fila in filas:
        celdas = t.add_row().cells
        for celda, texto in zip(celdas, fila):
            celda.text = ""
            corrida = celda.paragraphs[0].add_run(texto)
            corrida.font.size = Pt(9.5)
    documento.add_paragraph()


def imagen(documento: Document, ruta: pathlib.Path, pie: str) -> None:
    if not ruta.exists():
        return
    parrafo = documento.add_paragraph()
    parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    parrafo.add_run().add_picture(str(ruta), width=Inches(6.2))
    leyenda = documento.add_paragraph()
    leyenda.alignment = WD_ALIGN_PARAGRAPH.CENTER
    corrida = leyenda.add_run(pie)
    corrida.font.size = Pt(8.5)
    corrida.font.color.rgb = GRIS


# ---------------------------------------------------------------------------
# Documento general y técnico
# ---------------------------------------------------------------------------

def documento_tecnico() -> None:
    doc = Document()
    preparar(doc)
    portada(
        doc,
        "Documentación general y técnica",
        "Flash Campaign · creador de anuncios para Google Search, Google Display y Meta",
    )

    doc.add_heading("1. Descripción general", level=1)
    doc.add_paragraph(
        "Flash Campaign es un aplicativo web y webapp que produce propuestas de anuncios y "
        "campañas digitales de marca y performance para tres plataformas: Google Search, "
        "Google Display y Meta. Resuelve dos trabajos distintos con dos herramientas distintas: "
        "los anuncios de búsqueda se generan con un asistente de inteligencia artificial y las "
        "piezas gráficas se arman en un editor propio que replica un solo diseño a todos los formatos."
    )

    doc.add_heading("Problema que resuelve", level=2)
    vinetas(doc, [
        "Producir una campaña de display exige rehacer la misma pieza en 15 o 20 tamaños distintos.",
        "Redactar anuncios de búsqueda obliga a leer la landing, revisar la competencia y contar caracteres uno por uno.",
        "El material de marca (paleta, logotipos, tipografía) se dispersa entre carpetas y se aplica de forma inconsistente.",
    ])

    doc.add_heading("Objetivo", level=2)
    doc.add_paragraph(
        "Reducir a una sola pasada el trabajo de producir una campaña completa: un briefing para "
        "Search y un lienzo base de 300 × 250 para Display y Meta."
    )

    doc.add_heading("Usuarios", level=2)
    tabla(doc, ["Perfil", "Qué puede hacer"], [
        ["Administrador", "Todo lo del usuario, más crear y habilitar cuentas, ver las campañas y las imágenes de todo el equipo, editar la marca del sistema y revisar el estado del despliegue."],
        ["Usuario", "Crear y trabajar sus campañas, sus marcas y sus imágenes. No ve el material de otras cuentas."],
    ])

    doc.add_heading("Flujo principal", level=2)
    vinetas(doc, [
        "Entrar con correo y contraseña.",
        "Crear la campaña: nombre, plataformas (una, dos o las tres) y marca (Zurich, otra marca guardada o estilo libre).",
        "Search: completar el briefing y pedir la propuesta a Char B, que lee las páginas indicadas antes de redactar.",
        "Display y Meta: diseñar el lienzo base de 300 × 250, guardar y pulsar Replicar.",
        "Revisar formato por formato, ajustar los que lo necesiten y seleccionar los que se exportan.",
        "Exportar en JPG y/o HTML5 y cerrar la campaña como realizada.",
    ])

    doc.add_heading("Entradas y resultados", level=2)
    tabla(doc, ["Entrada", "Resultado"], [
        ["Briefing de Search: tipo de anuncios, CTA, gancho comercial, destino, URLs de referencia, fotos e indicaciones",
         "Resumen del producto y la competencia, grupos de anuncios con 15 títulos y 4 descripciones, rutas, palabras clave por concordancia, negativas, extensiones y recomendaciones. Exportable a CSV."],
        ["Lienzo base 300 × 250 con fondo, formas, textos, logotipo y CTA",
         "19 formatos de Google Display y Meta, con vista previa, selección por casilla y exportación en ZIP (JPG y HTML5 con clickTag)."],
    ])

    doc.add_heading("2. Alcance", level=1)
    doc.add_heading("Incluido", level=2)
    vinetas(doc, [
        "Acceso con usuario y contraseña, sesión firmada y roles de administrador y usuario.",
        "Creación de campañas para una, dos o las tres plataformas.",
        "Asistente Char B sobre la API de Google Gemini, con lectura real de las páginas indicadas.",
        "Editor gráfico con fondo de color o fotografía, filtros con intensidad, formas, cuadros de texto, logotipos y botón CTA.",
        "Arrastre libre con el mouse, redimensión proporcional desde las esquinas y libre desde los lados, y desplazamiento con teclado.",
        "Reglas de replicación con tres modos: escala uniforme, banda horizontal y columna vertical.",
        "Biblioteca de imágenes que lista directamente el bucket R2 de la cuenta, con subida propia y permisos por usuario.",
        "Marcas guardadas, creación de marcas y estilo libre.",
        "Proyectos en curso y realizados, historial con hora de creación y de exportación, y configuración de la cuenta.",
    ])

    doc.add_heading("Límites y procesos manuales", level=2)
    vinetas(doc, [
        "La tipografía Zurich Sans no es distribuible: la marca Zurich usa Arial, la alternativa que autoriza su brandbook.",
        "El aplicativo no publica campañas en Google Ads ni en Meta: entrega los archivos y los textos listos para cargar.",
        "La animación de banners HTML5 no está incluida: las piezas se exportan estáticas.",
        "Las fotografías con licencia de Envato deben descargarse con la cuenta del cliente (ver ENVATO_ASSETS.md).",
    ])

    doc.add_heading("3. Arquitectura", level=1)
    tabla(doc, ["Capa", "Tecnología", "Detalle"], [
        ["Frontend", "React 18 + TypeScript + Vite 6", "SPA con React Router. CSS propio con tokens de marca. Sin framework de UI de terceros."],
        ["Backend", "Cloudflare Pages Functions", "API en /api/*, ejecutada en el borde. Sin servidor propio que mantener."],
        ["Almacenamiento", "Cloudflare D1 (DB)", "Usuarios, marcas, proyectos, historial y metadatos de la biblioteca, en dos tablas que se crean solas: documentos y binarios. Admite también un espacio KV como alternativa."],
        ["Archivos", "Cloudflare R2 (IMAGENES)", "Biblioteca de imágenes: el catálogo ya cargado en el bucket y las subidas de los usuarios, que viven bajo el prefijo img/. Si R2 no está enlazado, se usa KV."],
        ["Inteligencia artificial", "Google Gemini", "Modelo configurable; por defecto gemini-2.5-flash con salida JSON validada por esquema."],
        ["Tipografía", "@fontsource/roboto", "Roboto 400/500/600 empaquetada con el aplicativo, sin llamadas externas."],
        ["Compresión", "fflate", "Generación del ZIP de exportación en el navegador."],
    ])

    doc.add_heading("Autenticación", level=2)
    doc.add_paragraph(
        "Las contraseñas se guardan como hash PBKDF2-SHA256 con sal y 150.000 iteraciones "
        "(shared/passwords.ts). La sesión viaja en una cookie HttpOnly, SameSite=Lax y Secure en "
        "producción, firmada con HMAC-SHA256 (shared/sesion.ts) y con 12 horas de vigencia. El "
        "mensaje de error es el mismo para usuario inexistente y contraseña incorrecta, de modo que "
        "no se puede enumerar cuentas."
    )

    doc.add_heading("Integraciones", level=2)
    vinetas(doc, [
        "Google Gemini: la clave vive en las variables de Cloudflare y nunca llega al navegador. Si el modelo configurado no está disponible, el servidor prueba alternativas antes de fallar.",
        "Lectura de páginas: el servidor descarga el HTML de las URLs del briefing, extrae título, meta descripción y texto, y lo entrega al modelo. Si una página no responde, se informa en pantalla y el modelo lo declara en sus recomendaciones.",
        "Biblioteca de Cloudflare: con el bucket R2 enlazado, el aplicativo recorre sus objetos y muestra las imágenes que ya estaban cargadas, usando las carpetas como etiquetas de búsqueda. Si el bucket tiene dominio público, MEDIA_BASE_URL hace que se sirvan desde ahí; si no, pasan por la API y exigen sesión.",
        "Detección de enlaces: si el espacio KV o el bucket no se llaman FLASH_KV y MEDIA, el aplicativo los identifica por su forma y muestra el nombre encontrado en Configuración.",
    ])

    doc.add_heading("4. Estructura del proyecto", level=1)
    tabla(doc, ["Ruta", "Función"], [
        ["data/usuarios.json", "Semilla de cuentas versionada en GitHub. Sólo hashes, nunca contraseñas en claro."],
        ["shared/passwords.ts", "Derivación y verificación de contraseñas con WebCrypto."],
        ["shared/sesion.ts", "Firma y lectura del token de sesión."],
        ["functions/_lib/", "Entorno y tipos del worker, capa de almacenamiento, usuarios e historial."],
        ["functions/api/", "Rutas de la API: auth, usuarios, marcas, proyectos, biblioteca, ia y estado."],
        ["src/lib/replicar.ts", "Reglas de replicación a los 19 formatos."],
        ["src/lib/render.ts", "Dibujado del banner en canvas: alimenta miniatura, editor y exportación JPG."],
        ["src/lib/exportar.ts", "Generación del HTML5 con clickTag y del ZIP de entrega."],
        ["src/lib/diseno.ts", "Fábricas de elementos y armado inicial de la campaña gráfica."],
        ["src/components/Lienzo.tsx", "Superficie de edición con arrastre, redimensión y teclado."],
        ["src/components/Inspector.tsx", "Panel de propiedades del fondo, del enlace y del elemento activo."],
        ["src/pages/", "Una pantalla por ruta."],
        ["public/brand/", "Logotipos de Flash Campaign, de Zurich y endoso de Gamonal."],
        ["scripts/", "Generación de la marca, semilla de usuarios y documentos."],
        ["scripts/qa/", "Recorrido funcional, prueba visual de replicación, capturas y doble de pruebas de Gemini."],
        ["Info-Zurich/", "Material de marca entregado con el encargo: brandbook y referencias de campañas."],
    ])

    doc.add_heading("5. Rutas del aplicativo", level=1)
    tabla(doc, ["Ruta", "Pantalla", "Acceso"], [
        ["/login", "Acceso", "Público"],
        ["/", "Inicio", "Con sesión"],
        ["/campanas/nueva", "Crear campaña", "Con sesión"],
        ["/campanas/:id/search", "Creador de Google Search con Char B", "Con sesión"],
        ["/campanas/:id/editor", "Editor de banners", "Con sesión"],
        ["/proyectos", "Proyectos en curso", "Con sesión"],
        ["/proyectos/realizados", "Proyectos realizados", "Con sesión"],
        ["/marcas", "Marcas", "Con sesión"],
        ["/biblioteca", "Biblioteca de imágenes", "Con sesión"],
        ["/historial", "Historial de trabajos", "Con sesión"],
        ["/configuracion", "Configuración y estado del despliegue", "Con sesión"],
        ["/usuarios", "Administración de cuentas", "Sólo administrador"],
    ])

    doc.add_heading("6. API", level=1)
    tabla(doc, ["Método y ruta", "Función"], [
        ["GET /api/estado", "Diagnóstico: qué está enlazado y qué falta configurar."],
        ["POST /api/auth/login", "Inicia sesión y entrega la cookie firmada."],
        ["POST /api/auth/logout", "Cierra la sesión."],
        ["GET /api/auth/sesion", "Devuelve el usuario de la sesión vigente."],
        ["POST /api/auth/perfil", "Actualiza nombre, correo de contacto y zona horaria."],
        ["POST /api/auth/clave", "Cambia la contraseña propia."],
        ["GET / POST /api/usuarios", "Lista y crea cuentas (sólo administrador)."],
        ["PATCH /api/usuarios/:id", "Cambia rol, estado o contraseña de una cuenta."],
        ["GET / POST /api/marcas, DELETE /api/marcas/:id", "Marcas guardadas."],
        ["GET / POST /api/proyectos, GET / DELETE /api/proyectos/:id", "Campañas."],
        ["GET / POST /api/biblioteca, GET / DELETE /api/biblioteca/:id", "Biblioteca de imágenes y subidas de los usuarios."],
        ["GET /api/biblioteca/objeto?clave=…", "Entrega una imagen del catálogo que ya existía en el bucket R2."],
        ["GET / POST /api/historial", "Bitácora de trabajos."],
        ["POST /api/ia/search", "Genera la propuesta de Search con Char B."],
    ])

    doc.add_heading("7. Reglas de replicación", level=1)
    doc.add_paragraph(
        "El lienzo base es el 300 × 250. Cada formato de destino se clasifica por su proporción y "
        "se resuelve con una de tres reglas. Después de replicar, cualquier formato puede editarse a "
        "mano: al hacerlo queda marcado como ajustado y una nueva replicación no lo pisa salvo que se "
        "confirme expresamente."
    )
    tabla(doc, ["Modo", "Cuándo se aplica", "Qué hace"], [
        ["Escala uniforme", "Proporción parecida a la base (por ejemplo 336 × 280, 1080 × 1080, 1200 × 628)",
         "Escala geométrica acotada y anclas conservadas: lo que estaba pegado a un borde sigue pegado a ese borde."],
        ["Banda horizontal", "Proporción al menos 1,9 veces más ancha que la base (728 × 90, 970 × 250, 468 × 60)",
         "El orden de lectura vertical se convierte en horizontal: arriba pasa a la izquierda, centro al centro y abajo a la derecha. La decoración de fondo escala por alto."],
        ["Columna vertical", "Proporción al menos la mitad de ancha que la base (160 × 600, 300 × 1050, 1080 × 1920)",
         "Se escala por ancho, se agrupan los elementos en filas y el aire sobrante se reparte entre ellas sin deformar nada. La decoración escala por ancho."],
    ])
    doc.add_paragraph(
        "Una forma o imagen que sobresale más del 12 % del lienzo se trata como recurso gráfico de "
        "fondo y se reubica por proporción sin recortarse; los cuadros de texto nunca se consideran "
        "decoración, porque siempre deben poder leerse."
    )

    doc.add_heading("8. Exportación", level=1)
    vinetas(doc, [
        "JPG por pieza, con calidad configurable entre 60 % y 100 %.",
        "HTML5 autónomo por pieza, con <meta name=\"ad.size\"> y variable clickTag, tal como exigen Google Ads y Campaign Manager.",
        "Todo se entrega en un ZIP con carpetas jpg/ y html/ y un LEEME.txt que lista el contenido.",
        "La selección de formatos se hace con casillas: se exporta lo marcado, no siempre todo.",
    ])

    doc.add_heading("9. Requisitos, instalación y despliegue", level=1)
    doc.add_paragraph("Requisitos: Node.js 20 o superior y una cuenta de Cloudflare con Pages habilitado.")
    tabla(doc, ["Comando", "Qué hace"], [
        ["npm install", "Instala las dependencias."],
        ["npm run dev:cf", "Compila y levanta el aplicativo completo en http://127.0.0.1:8788."],
        ["npm run build", "Comprueba tipos del frontend y de las funciones y compila a dist/."],
        ["npm run typecheck", "Sólo comprobación de tipos."],
        ["npm run deploy", "Compila y publica en Cloudflare Pages."],
        ['node scripts/seed_users.mjs "<clave>"', "Regenera data/usuarios.json con el hash de la contraseña del administrador."],
    ])

    doc.add_heading("Variables y enlaces de Cloudflare", level=2)
    tabla(doc, ["Nombre", "Tipo", "Obligatorio", "Para qué sirve"], [
        ["DB", "Base de datos D1", "Sí", "Usuarios, marcas, proyectos e historial."],
        ["IMAGENES", "Bucket R2", "Recomendado", "Biblioteca de imágenes: catálogo existente y subidas de los usuarios."],
        ["JWT_SECRET o SESSION_SECRET", "Secreto", "Recomendado", "Firma de la cookie de sesión."],
        ["GEMINI_API_KEY", "Secreto", "Sí para Search", "Clave de Google Gemini que usa Char B."],
        ["GEMINI_MODEL", "Variable", "No", "Modelo a usar. Por defecto gemini-2.5-flash."],
        ["GEMINI_BASE_URL", "Variable", "No", "Base alternativa de la API (proxy o entorno de prueba)."],
        ["MEDIA_BASE_URL", "Variable", "No", "Dominio público del bucket: si se define, las imágenes se sirven desde ahí."],
        ["MEDIA_MANIFEST_URL", "Variable", "No", "Manifiesto JSON alternativo, sólo si el catálogo no vive en el bucket enlazado."],
        ["D1_BINDING / R2_BINDING / KV_BINDING", "Variable", "No", "Nombre real del enlace si usa otra etiqueta."],
    ])
    doc.add_paragraph(
        "Sin FLASH_KV el aplicativo funciona pero avisa en pantalla que el almacenamiento es temporal. "
        "La sección Configuración muestra en todo momento el estado real de cada enlace y el nombre con el "
        "que lo encontró."
    )
    doc.add_paragraph(
        "El repositorio no incluye un archivo wrangler.toml de forma deliberada. Cuando un proyecto de "
        "Cloudflare Pages tiene ese archivo, la plataforma lo toma como fuente de verdad e ignora los "
        "enlaces y variables definidos en el panel; como aquí la configuración vive en el panel, agregarlo "
        "dejaría el aplicativo sin KV, sin R2 y sin la clave de Gemini."
    )

    doc.add_heading("10. Seguridad y privacidad", level=1)
    vinetas(doc, [
        "Contraseñas con PBKDF2-SHA256, sal por usuario y 150.000 iteraciones. Comparación en tiempo constante.",
        "Cookie de sesión HttpOnly, SameSite=Lax y Secure sobre HTTPS.",
        "La clave de Gemini sólo existe en el servidor; el navegador nunca la ve.",
        "Cada usuario ve únicamente sus campañas y sus imágenes; el administrador ve todo.",
        "La biblioteca no es pública: los archivos se sirven a través de la API y exigen sesión.",
        "Cabeceras X-Content-Type-Options, X-Frame-Options, Referrer-Policy y Permissions-Policy definidas en public/_headers.",
        "El administrador no puede quitarse a sí mismo el acceso, de modo que siempre queda una cuenta con control.",
    ])

    doc.add_heading("11. Pruebas ejecutadas", level=1)
    tabla(doc, ["Prueba", "Cómo se ejecutó", "Resultado"], [
        ["Tipos", "npm run typecheck sobre frontend y funciones", "Sin errores."],
        ["API", "curl contra el aplicativo levantado: login correcto e incorrecto, sesión, proyectos, marcas, biblioteca, usuarios e historial", "Todas las rutas responden como se espera; sin sesión devuelven 401."],
        ["Recorrido funcional", "scripts/qa/recorrido.mjs con navegador Chromium real", "27 comprobaciones en verde: login, campaña, Char B, edición, arrastre, redimensión, replicación, vista previa, exportación ZIP y CSV, alta de usuario y responsive."],
        ["Char B", "scripts/qa/stub-gemini.mjs como doble de pruebas de la API de Gemini", "Se comprobó que el servidor lee la página del briefing, arma el prompt y recorta los textos a los límites de Google Ads."],
        ["Replicación", "scripts/qa/replicacion.mjs con un lienzo base al estilo de las referencias de Zurich", "Los 19 formatos se generan legibles y sin desbordes."],
        ["Responsive", "Chromium en 390 × 844, 834 × 1112 y 1600 × 1000", "Sin scroll horizontal en ninguno de los tres anchos."],
    ])

    doc.add_heading("12. Capturas del producto", level=1)
    for nombre, pie in [
        ("pantalla_inicio_flash_campaign.jpg", "Inicio: apertura de la tarea principal, estado de la cuenta y últimos movimientos."),
        ("pantalla_search_flash_campaign.jpg", "Creador de Google Search: briefing a la izquierda y propuesta de Char B con contador de caracteres a la derecha."),
        ("pantalla_editor_flash_campaign.jpg", "Editor de banners: lista de formatos, lienzo base y panel de propiedades."),
        ("pantalla_formatos_flash_campaign.jpg", "Vista previa de los 19 formatos generados desde el lienzo base."),
    ]:
        imagen(doc, CAPTURAS / nombre, pie)
        doc.add_paragraph()

    doc.add_heading("13. Decisiones técnicas y riesgos", level=1)
    tabla(doc, ["Decisión", "Alternativa evaluada", "Motivo"], [
        ["Cloudflare Pages Functions para la API", "Servidor Node propio", "El proyecto ya vive en Cloudflare; no agrega infraestructura que mantener y la clave de IA queda del lado servidor."],
        ["D1 como almacén principal", "KV", "El proyecto en Cloudflare ya tenía enlazada una base D1, así que el aplicativo se adaptó a ella en vez de exigir un recurso nuevo. Los documentos se guardan en una tabla clave-valor y KV queda como alternativa."],
        ["Listar el bucket R2 en vivo", "Mantener un manifiesto JSON del catálogo", "El catálogo no se desincroniza: lo que hay en el bucket es lo que se ve, sin un archivo intermedio que actualizar."],
        ["Enlaces por nombre explícito, con KV_BINDING y R2_BINDING como escape", "Detectar el enlace por su forma", "Pages expone enlaces propios como ASSETS que responden a cualquier propiedad: una búsqueda por forma llegó a elegir ASSETS como bucket y tumbó la biblioteca. El nombre explícito no se equivoca."],
        ["Canvas para dibujar el banner", "DOM con divs posicionados", "Lo que se ve en el editor es exactamente lo que se exporta en JPG, sin diferencias de renderizado."],
        ["Salida estructurada de Gemini con esquema", "Pedir texto libre y parsearlo", "Evita respuestas mal formadas y permite validar los límites de caracteres del lado servidor."],
        ["Roboto empaquetada", "Google Fonts por CDN", "El aplicativo no depende de un tercero para renderizar su interfaz."],
        ["Imágenes incrustadas como data URI en el HTML exportado", "Referenciar la ruta del aplicativo", "La pieza se sube a Google Ads como archivo suelto: no puede depender de rutas del servidor de origen."],
    ])

    doc.add_heading("Riesgos y limitaciones conocidas", level=2)
    vinetas(doc, [
        "Sin GEMINI_API_KEY el creador de Search no genera; el editor de Display y Meta funciona igual.",
        "Sin FLASH_KV los datos no persisten entre reinicios del worker.",
        "El historial guarda los últimos 500 movimientos: es una bitácora de trabajo, no un registro de auditoría permanente.",
        "La lectura de páginas depende de que la landing entregue su contenido en HTML; una página que se arma sólo con JavaScript entrega menos contexto al modelo.",
    ])

    doc.add_heading("14. Mantenimiento y pendientes", level=1)
    vinetas(doc, [
        "Respaldo: exportar el contenido del espacio KV con wrangler kv key list / get antes de cambios mayores.",
        "Pendiente del cliente: descargar las dos fotografías de Envato registradas en ENVATO_ASSETS.md y dejarlas en public/media/.",
        "Las tablas de D1 se crean solas en la primera petición: no hay migraciones que ejecutar a mano.",
    ])

    doc.add_heading("15. Historial de cambios", level=1)
    tabla(doc, ["Versión", "Fecha", "Cambio"], [
        ["1.0.0", FECHA.strftime("%d-%m-%Y"), "Primera entrega completa: acceso, campañas, Char B, editor con replicación, biblioteca, marcas, historial, configuración y administración de usuarios."],
        ["1.1.0", FECHA.strftime("%d-%m-%Y"), "La biblioteca lista directamente el bucket R2 de la cuenta y el HTML exportado incrusta sus imágenes."],
        ["1.2.0", FECHA.strftime("%d-%m-%Y"), "El almacenamiento se adapta al proyecto real de Cloudflare: base D1 enlazada como DB, bucket R2 como IMAGENES y secreto de sesión como JWT_SECRET."],
    ])

    ruta = DOCS / "documentacion_general_tecnica_flash_campaign.docx"
    doc.save(str(ruta))
    print("✓", ruta)


# ---------------------------------------------------------------------------
# Documento publicitario
# ---------------------------------------------------------------------------

def documento_publicitario() -> None:
    doc = Document()
    preparar(doc)
    portada(doc, "Descripción publicitaria", "Flash Campaign · textos comerciales y mensajes clave")

    doc.add_heading("Ficha", level=1)
    tabla(doc, ["Campo", "Contenido"], [
        ["Nombre comercial", "Flash Campaign"],
        ["Nombre normalizado", "flash_campaign"],
        ["Categoría", "Herramienta de producción publicitaria digital"],
        ["Pensado y creado por", "Gamonal"],
        ["Plataformas que cubre", "Google Search, Google Display y Meta"],
        ["Acceso", "Web y webapp, con cuentas por usuario"],
    ])

    doc.add_heading("Descripción de una línea", level=1)
    doc.add_paragraph("Un lienzo, toda la campaña: Flash Campaign convierte un banner de 300 × 250 en 19 formatos y redacta los anuncios de búsqueda con inteligencia artificial.")

    doc.add_heading("Descripción breve para presentación", level=1)
    doc.add_paragraph(
        "Flash Campaign es el aplicativo con el que un equipo de marketing arma una campaña digital "
        "completa en una sola pasada. Se diseña una vez el banner base y el sistema genera los catorce "
        "formatos de Google Display y los cinco de Meta con reglas propias para piezas horizontales y "
        "verticales. En paralelo, el asistente Char B lee la página de la promoción y propone los "
        "títulos, descripciones y palabras clave de Google Search ya ajustados a los límites de la plataforma."
    )

    doc.add_heading("Descripción completa para sitio web", level=1)
    doc.add_paragraph(
        "Producir una campaña digital significa repetir el mismo trabajo muchas veces: el mismo mensaje "
        "en quince tamaños de banner, los mismos beneficios recortados a treinta caracteres, la misma "
        "paleta buscada otra vez en el manual de marca. Flash Campaign concentra ese trabajo en un solo lugar."
    )
    doc.add_paragraph(
        "Para Google Display y Meta hay un editor propio: fondo de color o fotografía con filtros de "
        "intensidad regulable, formas con colores de marca, cuadros de texto con control de tipografía y "
        "márgenes, logotipos y botón de llamada a la acción. Todo se mueve con el mouse y se redimensiona "
        "en proporción desde las esquinas. Cuando el lienzo base de 300 × 250 está listo, un botón replica "
        "el diseño a los diecinueve formatos y cada uno queda editable por separado."
    )
    doc.add_paragraph(
        "Para Google Search está Char B, el asistente de inteligencia artificial del aplicativo. Recibe el "
        "briefing, entra a la página de la promoción, entiende el producto, la oferta y la competencia, y "
        "devuelve grupos de anuncios con quince títulos, cuatro descripciones, rutas de visualización, "
        "palabras clave separadas por concordancia amplia, de frase y exacta, palabras negativas y "
        "extensiones. El aplicativo cuenta los caracteres y avisa si algo se pasa del límite."
    )
    doc.add_paragraph(
        "La marca Zurich viene creada con su paleta, sus logotipos y su tipografía. Cada equipo puede "
        "crear las suyas o trabajar en estilo libre. Las imágenes viven en una biblioteca compartida donde "
        "cada persona ve las suyas y el administrador ve todo. Y cada exportación queda registrada en el "
        "historial con su fecha y su hora."
    )

    doc.add_heading("Descripción para ficha de producto", level=1)
    doc.add_paragraph(
        "Aplicativo web para producir campañas de Google Search, Google Display y Meta. Editor gráfico con "
        "replicación automática a 19 formatos, asistente de IA para anuncios de búsqueda, biblioteca de "
        "imágenes, marcas guardadas e historial de trabajos. Exportación en JPG y HTML5 con clickTag."
    )

    doc.add_heading("Problema, solución y público", level=1)
    tabla(doc, ["Bloque", "Contenido"], [
        ["Problema", "Cada campaña obliga a rehacer la misma pieza en quince o veinte tamaños y a redactar decenas de textos contando caracteres a mano."],
        ["Solución", "Un lienzo base que se replica con reglas pensadas para cada proporción, y un asistente que lee la promoción antes de escribir."],
        ["Contexto", "Equipos de marketing y agencias que trabajan campañas de performance con calendario apretado y varias marcas a la vez."],
        ["Público objetivo", "Responsables de medios digitales, diseñadores de performance y equipos internos de marketing."],
    ])

    doc.add_heading("Propuesta de valor", level=1)
    vinetas(doc, [
        "Diseñas una vez y obtienes toda la campaña.",
        "Los textos de búsqueda nacen leyendo la landing real, no una descripción de segunda mano.",
        "La marca se aplica sola: paleta, logotipos y tipografía ya cargados.",
        "Lo que ves en el editor es exactamente lo que se exporta.",
    ])

    doc.add_heading("Atributos y beneficios verificables", level=1)
    tabla(doc, ["Capacidad real", "Beneficio"], [
        ["19 formatos generados desde un lienzo de 300 × 250", "Se elimina la parte repetitiva de la producción gráfica."],
        ["Tres reglas de replicación según proporción", "Las piezas horizontales y verticales quedan armadas, no sólo encogidas."],
        ["Cada formato editable y marcado cuando se ajusta a mano", "Se puede afinar una pieza sin miedo a perder el trabajo en la siguiente replicación."],
        ["Char B lee las URLs del briefing antes de redactar", "Los anuncios hablan del producto real, con sus precios y coberturas."],
        ["Contador de caracteres por título y descripción", "Nada se rechaza al cargarlo en Google Ads."],
        ["Exportación en JPG y HTML5 con clickTag y ad.size", "Los archivos entran directo en Google Ads y Campaign Manager."],
        ["Biblioteca con permisos por usuario", "Cada persona trabaja con su material sin mezclarlo con el de otros."],
        ["Historial con hora de creación y exportación", "Queda trazabilidad de qué se produjo y cuándo."],
    ])

    doc.add_heading("Casos de uso", level=1)
    vinetas(doc, [
        "Lanzamiento de una promoción de seguro de auto con anuncios de búsqueda y banners para toda la red de Display.",
        "Adaptación de una campaña existente a Meta sin volver a diseñarla.",
        "Producción simultánea de varias marcas desde un mismo equipo.",
        "Renovación mensual de una oferta: se cambia el gancho, se replica y se exporta.",
    ])

    doc.add_heading("Mensajes clave", level=1)
    vinetas(doc, [
        "Un lienzo, toda la campaña.",
        "Los anuncios nacen leyendo tu propia promoción.",
        "De 300 × 250 a diecinueve formatos, sin rehacer el diseño.",
        "Listo para Google Ads: JPG y HTML5 con clickTag.",
    ])

    doc.add_heading("Texto para LinkedIn", level=1)
    doc.add_paragraph(
        "Producir una campaña digital tiene una parte creativa y otra que es pura repetición: el mismo "
        "mensaje en quince tamaños, los mismos beneficios recortados a treinta caracteres.\n\n"
        "Flash Campaign se ocupa de la segunda parte. Diseñas el banner de 300 × 250 y el aplicativo "
        "genera los catorce formatos de Google Display y los cinco de Meta, con reglas distintas para las "
        "piezas horizontales y las verticales. En paralelo, Char B entra a la página de tu promoción, "
        "entiende el producto y la oferta, y propone títulos, descripciones y palabras clave por "
        "concordancia, con los caracteres ya contados.\n\n"
        "Exporta en JPG y en HTML5 con clickTag, listo para cargar. Pensado y creado por Gamonal."
    )

    doc.add_heading("Textos breves para redes sociales", level=1)
    vinetas(doc, [
        "Diseña el 300 × 250. Nosotros hacemos los otros 18 formatos. #FlashCampaign",
        "Char B lee tu landing y escribe los anuncios de búsqueda. Con los caracteres ya contados.",
        "De un lienzo a toda la campaña: Google Search, Google Display y Meta.",
    ])

    doc.add_heading("Palabras clave", level=1)
    doc.add_paragraph(
        "producción de banners, replicación de formatos, Google Display, Google Search, Meta Ads, "
        "anuncios responsivos, palabras clave por concordancia, clickTag, HTML5 para Google Ads, "
        "editor de banners, campañas de performance, marketing digital"
    )

    doc.add_heading("Nota sobre el uso comercial de estos textos", level=1)
    doc.add_paragraph(
        "Todos los argumentos anteriores describen capacidades que están implementadas y probadas en la "
        "versión 1.0.0. No se incluyen cifras de rendimiento, testimonios ni comparaciones con otras "
        "herramientas porque no existe medición que los respalde."
    )

    ruta = DOCS / "descripcion_publicitaria_flash_campaign.docx"
    doc.save(str(ruta))
    print("✓", ruta)


if __name__ == "__main__":
    documento_tecnico()
    documento_publicitario()
